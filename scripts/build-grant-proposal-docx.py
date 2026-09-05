#!/usr/bin/env python3
"""Build the Google Docs-targeted grant proposal DOCX from its Markdown source."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


INK = RGBColor(0x00, 0x00, 0x00)
MUTED = RGBColor(0x55, 0x55, 0x55)
BORDER = "DADCE0"
INLINE_RE = re.compile(r"(\[[^\]]+\]\(https?://[^\s)]+\)|\*\*.+?\*\*|`.+?`|https?://[^\s]+)")


def set_font(run, name: str, size: float, color: RGBColor = INK, bold=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_style_font(style, name: str, size: float, color: RGBColor, bold: bool) -> None:
    style.font.name = name
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "22")
    properties.extend([fonts, color, underline, size])
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def normalize_text(text: str) -> str:
    return (
        text.replace("—", "-")
        .replace("–", "-")
        .replace("→", "->")
        .replace("“", '"')
        .replace("”", '"')
        .replace("’", "'")
    )


def add_inline(paragraph, text: str) -> None:
    text = normalize_text(text)
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            set_font(paragraph.add_run(text[cursor : match.start()]), "Arial", 11)
        token = match.group(0)
        if token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        elif token.startswith("**"):
            content = token[2:-2]
            run = paragraph.add_run(content)
            set_font(run, "Arial", 11, bold=True)
            if content.startswith("SUBMISSION INPUT REQUIRED:"):
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, "Courier New", 10.5)
        else:
            trimmed = token.rstrip(".,;:")
            suffix = token[len(trimmed) :]
            add_hyperlink(paragraph, trimmed, trimmed)
            if suffix:
                set_font(paragraph.add_run(suffix), "Arial", 11)
        cursor = match.end()
    if cursor < len(text):
        set_font(paragraph.add_run(text[cursor:]), "Arial", 11)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = tc_mar.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            tc_mar.append(tag)
        tag.set(qn("w:w"), str(value))
        tag.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table_pr = table._tbl.tblPr
    table_w = table_pr.first_child_found_in("w:tblW")
    if table_w is None:
        table_w = OxmlElement("w:tblW")
        table_pr.append(table_w)
    table_w.set(qn("w:w"), str(sum(widths)))
    table_w.set(qn("w:type"), "dxa")
    table_ind = table_pr.first_child_found_in("w:tblInd")
    if table_ind is None:
        table_ind = OxmlElement("w:tblInd")
        table_pr.append(table_ind)
    table_ind.set(qn("w:w"), "0")
    table_ind.set(qn("w:type"), "dxa")
    layout = table_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_quiet_borders(table) -> None:
    table_pr = table._tbl.tblPr
    borders = table_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), BORDER)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_numbering(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "276")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, indent, spacing])
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    r_pr.append(fonts)
    level.extend([start, num_fmt, level_text, justification, p_pr, r_pr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_number(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_element])
    p_pr.append(num_pr)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    set_style_font(normal, "Arial", 11, INK, False)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in (
        ("Heading 1", 20, INK, 20, 6),
        ("Heading 2", 16, INK, 18, 6),
        ("Heading 3", 14, RGBColor(0x43, 0x43, 0x43), 16, 4),
    ):
        style = document.styles[name]
        set_style_font(style, "Arial", size, color, False)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    component = document.styles.add_style("Component Heading", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(component, "Arial", 12, INK, True)
    component.paragraph_format.space_before = Pt(12)
    component.paragraph_format.space_after = Pt(4)
    component.paragraph_format.line_spacing = 1.0
    component.paragraph_format.keep_with_next = True


def add_table(document: Document, lines: list[str]) -> None:
    rows = [[cell.strip() for cell in line.strip().strip("|").split("|")] for line in lines]
    rows = [row for index, row in enumerate(rows) if index != 1]
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    widths = [504, 3024, 4752, 1080]
    set_table_geometry(table, widths)
    set_quiet_borders(table)
    repeat_table_header(table.rows[0])

    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            paragraph = table.cell(row_index, column_index).paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            if column_index in (0, 3):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(paragraph, value)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
            if row_index == len(rows) - 1:
                for run in paragraph.runs:
                    run.bold = True
    document.add_paragraph()


def build(markdown_path: Path, output_path: Path) -> None:
    document = Document()
    configure_document(document)
    document.core_properties.title = "Hedgents Stablecoin Rail - Developer Tooling Grant Proposal"
    document.core_properties.subject = "Solana Foundation Developer Tooling Grant Proposal"
    document.core_properties.author = "Hedgents"
    num_id = add_numbering(document)

    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    index = 0
    first_title = True
    while index < len(lines):
        raw = lines[index].rstrip()
        if not raw:
            index += 1
            continue
        if raw.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(document, table_lines)
            continue
        if raw.startswith("# ") and first_title:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(3)
            run = paragraph.add_run(normalize_text(raw[2:]))
            set_font(run, "Arial", 26, INK, bold=False)
            first_title = False
        elif raw.startswith("#### "):
            document.add_paragraph(normalize_text(raw[5:]), style="Component Heading")
        elif raw.startswith("### "):
            document.add_paragraph(normalize_text(raw[4:]), style="Heading 3")
        elif raw.startswith("## "):
            document.add_paragraph(normalize_text(raw[3:]), style="Heading 2")
        elif re.match(r"^\d+\. ", raw):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.15
            add_inline(paragraph, re.sub(r"^\d+\. ", "", raw))
            apply_number(paragraph, num_id)
        else:
            paragraph = document.add_paragraph()
            if raw == "Solana Foundation application":
                paragraph.paragraph_format.space_after = Pt(16)
                run = paragraph.add_run(raw)
                set_font(run, "Arial", 11, MUTED)
            else:
                if raw.startswith("**Beta scope:**"):
                    paragraph.paragraph_format.keep_together = True
                add_inline(paragraph, raw)
        index += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("markdown", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build(args.markdown, args.output)


if __name__ == "__main__":
    main()
